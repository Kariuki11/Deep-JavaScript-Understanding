from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import cairosvg
import os

HTML_PATH = "conceptual_framework.html"
DOCX_PATH = "conceptual_framework.docx"
SVG_PATH = "diagram.svg"
PNG_PATH = "diagram.png"

# Read HTML
with open(HTML_PATH, "r", encoding="utf-8") as f:
    html = f.read()

soup = BeautifulSoup(html, "lxml")

# Try to capture a page title
title = (soup.title.string.strip() if soup.title and soup.title.string else None)

# Extract the inline SVG markup
svg_tag = soup.find("svg")
if not svg_tag:
    raise SystemExit("No <svg> found in the HTML. Nothing to convert.")

# ✅ Ensure SVG has dimensions (fix for CairoSVG)
if not svg_tag.has_attr("width"):
    svg_tag["width"] = "1200px"
if not svg_tag.has_attr("height"):
    svg_tag["height"] = "700px"

# Save SVG to file
with open(SVG_PATH, "w", encoding="utf-8") as f:
    f.write(str(svg_tag))

# Create Word doc (A4 landscape)
doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11.69)   # A4 landscape
section.page_height = Inches(8.27)
for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(section, attr, Inches(0.79))  # ~20mm

# Optionally add title from <title>
if title:
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(16)
    run.bold = True

try:
    # 🚀 Try inserting the SVG directly (best clarity, Word 2016+)
    doc.add_picture(SVG_PATH, width=Inches(section.page_width.inches - section.left_margin.inches - section.right_margin.inches))
    print("Inserted SVG directly into Word ✅")
except Exception:
    # 🔄 Fallback: Convert to high-res PNG for older Word versions
    print("⚠️ SVG not supported, falling back to high-res PNG...")
    cairosvg.svg2png(url=SVG_PATH, write_to=PNG_PATH, scale=5.0, dpi=300)
    doc.add_picture(PNG_PATH, width=Inches(section.page_width.inches - section.left_margin.inches - section.right_margin.inches))

doc.save(DOCX_PATH)

# Clean up intermediates (optional)
try:
    os.remove(SVG_PATH)
    if os.path.exists(PNG_PATH):
        os.remove(PNG_PATH)
except Exception:
    pass

print(f"✅ Created {DOCX_PATH}")
